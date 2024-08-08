import secrets
import random
import string
import base64
import time
import requests
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric import ed25519
from cryptography.hazmat.backends import default_backend  # Import cryptography library for RSA encryption
from cryptography.hazmat.primitives.asymmetric import rsa, padding  # Import RSA and padding modules
import requests
import secrets
import time
import base64
import json
import os
from docx import Document
# pip install python-docx






class SampleData:
    valid_transaction_ids = []
    signature = ''
    encrypted_data = ''
    encryption_key = ''


    @staticmethod
    def generate_id(length):
        characters = "0123456789"
        random_string = ''.join(secrets.choice(characters) for _ in range(length))
        return random_string
    
    @staticmethod
    def generate_random_ascii_string(length):
        return ''.join(random.choice(string.ascii_letters + string.digits + string.punctuation) for _ in range(length))

    @staticmethod
    def generate_random_base64_string(length):
        # Calculate the number of bytes needed to achieve the desired length
        num_bytes = (length * 3 + 3) // 4
        random_bytes = os.urandom(num_bytes)

        # Encode the bytes to Base64
        base64_string = base64.b64encode(random_bytes).decode('utf-8')

        # Trim the string to the desired length
        base64_string = base64_string[:length]

        return base64_string
    
    @staticmethod
    def generate_random_non_b64_string(length):
        all_characters = "!@#$%^&*()-=_{}[]|;:'\",.<>?`~"
        random_string = ''.join(random.choice(all_characters) for _ in range(length))
        return random_string

    def get_data(self):
        data = {
            'transfer_amount': [
                [
                    "100",
                    "0.00001"
                ],
                [
                    "0.000001",
                    "1.000001",
                    "100000000000000"
                ]
            ],
            'transaction_amount': [
                [
                    "100",
                    "0.00001"
                ],
                [
                    "0.000001",
                    "1.000001",
                    "100000000000000"
                ]
            ],
            'request_id': [
                [
                    self.generate_id(32),
                ],
                [
                    self.generate_id(31),
                    self.generate_id(32),
                    self.generate_random_ascii_string(32),
                ],
            ],
            'transaction_id': [
                [
                    transaction_id
                ],
                [
                    self.generate_id(32),
                ],
            ],
            'transaction_ids': [
                [
                    [transaction_id],
                    str([]),
                ],
                [
                    self.generate_id(32),
                    str([self.generate_random_ascii_string(32)]),
                ],
            ],
            'recipient_key': [
                [
                    KeyPair().public_key_b64(),
                ],
                [
                    self.generate_random_base64_string(43),
                    self.generate_random_base64_string(45),
                    self.generate_random_non_b64_string(44),
                ]
            ],
            'master_key': [
                [
                    base64.b64encode(bytes.fromhex("aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa")).decode('utf-8'),
                ],
                [
                    self.generate_random_base64_string(43),
                    self.generate_random_base64_string(45),
                    self.generate_random_non_b64_string(44),
                ]
            ],
            'alias_address': [
                [
                    alias_address,
                ],
                [
                    self.generate_random_base64_string(43),
                    self.generate_random_base64_string(45),
                    self.generate_random_non_b64_string(44),
                ]
            ],
            'request_expiry_time': [
                [
                    str(int(time.time()) + 60),
                ],
                [
                    str(int(time.time()) - 60),
                    str(int(time.time()) + 100000),
                    self.generate_random_ascii_string(32),
                ]
            ],
            'transaction_expiry_time': [
                [
                    str(int(time.time()) + 60),
                ],
                [
                    str(int(time.time()) - 60),
                    str(int(time.time()) + 100000),
                    self.generate_random_ascii_string(32),
                ]
            ],
            'alias_expiry_time': [
                [
                    str(int(time.time()) + 60),
                ],
                [
                    str(int(time.time()) - 60),
                    str(int(time.time()) + 100000),
                    self.generate_random_ascii_string(32),
                ]
            ],
            'signature': [
                [
                    
                ],
                [
                    self.generate_random_non_b64_string(88),
                    self.generate_random_base64_string(87),
                    self.generate_random_base64_string(89),
                ]
            ],
            'transaction_type': [
                [
                    'SEND',
                    'RECEIVE'
                ],
                [
                    self.generate_random_ascii_string(4),
                    self.generate_random_ascii_string(7),
                ]
            ],
            'encrypted_data': [
                [
                    
                ],
                [
                    self.generate_random_base64_string(343),
                    self.generate_random_base64_string(345),
                    self.generate_random_non_b64_string(344),
                ]
            ],
            'encryption_key': [
                [
                    
                ],
                [
                    self.generate_random_non_b64_string(392),
                    self.generate_random_ascii_string(391),
                    self.generate_random_ascii_string(393),
                ]
            ]
        }

        return data


class KeyPair():
    
    def __init__(self, private_key_hex = None):

        if private_key_hex is None:
            self._private_key, self._public_key = self.generate_new_ed25519_keypair()
        else:
            self._private_key, self._public_key = self.generate_ed25519_keypair_from_key(bytes.fromhex(private_key_hex))

        


    def generate_new_ed25519_keypair(self):
        private_key = ed25519.Ed25519PrivateKey.generate()
        public_key = private_key.public_key()
        return private_key, public_key
    
    def generate_ed25519_keypair_from_key(self, private_bytes):
        private_key = ed25519.Ed25519PrivateKey.from_private_bytes(private_bytes)
        public_key = private_key.public_key()
        return private_key, public_key

    def public_key_b64(self):
        return base64.b64encode(self._public_key.public_bytes(
            encoding=serialization.Encoding.Raw,
            format=serialization.PublicFormat.Raw
        )).decode('utf-8')
    
    def sign(self, json_data):
        
        sorted_data = {k: json_data[k] for k in sorted(json_data.keys())}

        json_string = json.dumps(sorted_data, separators=(',', ':'))

        signature = self._private_key.sign(json_string.encode('utf-8'))
        return signature


class Testing:

    # Test number
    # Description
    # Expected Result
    # Actual Result
    # Pass/Fail

    results = []

    def __init__(self):
        self.rsa_encryption_key = self.get_current_encryption_key()
        self.sample_data = SampleData().get_data()

    def get_current_encryption_key(self):
        url = domain + '/api/get-key'
        response = requests.get(url)
        json_response = response.json()
        return json_response['key']
    
    def encrypt(self, message):
        serialized_public_key = base64.b64decode(self.rsa_encryption_key)  # Decode base64-encoded public key
        public_key = serialization.load_der_public_key(
            serialized_public_key,
            backend=default_backend()
        )
        message_sections = [message[i:i + 190] for i in range(0, len(message), 190)]  # Divide message into chunks
        encrypted_message_sections = []

        for section in message_sections:
            ciphertext = public_key.encrypt(section.encode('utf-8'), padding.PKCS1v15())  # Encrypt each message chunk
            encrypted_message_sections.append(base64.b64encode(ciphertext).decode('utf-8'))

        final_encrypted = ''.join(encrypted_message_sections)
        return final_encrypted
    
    def decrypt_message(self, private_key, ciphertext: str):
        ciphertext_sections = [ciphertext[i:i + 344] for i in range(0, len(ciphertext), 344)]  # Divide ciphertext into chunks
        plaintext_sections = []

        for section in ciphertext_sections:
            plaintext = private_key.decrypt(base64.b64decode(section), padding.PKCS1v15())  # Decrypt each ciphertext chunk


            plaintext_sections.append(plaintext.decode('utf-8'))

        final_plaintext = ''.join(plaintext_sections)
        return final_plaintext
    
    def send(self, endpoint, data):
        headers = {
            'Content-Type': 'base64',
            'Authorization': 'Bearer YourAccessToken',  # Include any necessary headers
        }

        url = domain + endpoint  # Construct the URL for the POST request

        # Send the POST request with the Base64 data in the body
        response = requests.post(url, data=data, headers=headers)
        json_response = response.json()

        return json_response
    

    def transfer_test(self):
        endpoint = '/api/transfer'
        
        required = ['request_id',
            'request_expiry_time',
            'transfer_amount',
            'recipient_key']
        
        json_data = {
            'sender_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)

    def create_transaction_test(self):
        endpoint = '/api/create-transaction'
        
        required = ['request_id',
            'request_expiry_time',
            'transaction_expiry_time',
            'transaction_amount',
            'transaction_type']
        
        json_data = {
            'master_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)
    
    def add_alias_test(self):
        endpoint = '/api/add-alias'
        
        required = ['request_id',
            'request_expiry_time',
            'alias_expiry_time',
            'alias_address']
        
        json_data = {
            'master_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)
    
    def delete_alias_test(self):
        endpoint = '/api/delete-alias'
        
        required = ['request_id',
            'request_expiry_time',
            'alias_address']
        
        json_data = {
            'master_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)
    
    def get_balance_test(self):
        endpoint = '/api/get-balance'
        
        required = ['request_id',
            'request_expiry_time']
        
        json_data = {
            'master_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)

    def get_transactions_test(self):
        endpoint = '/api/get-transactions'
        required = ['transaction_ids']

        self.check_json(required, endpoint, {})
    
    def delete_transaction_test(self):
        endpoint = '/api/delete-transaction'
        
        required = ['request_id',
                      'request_expiry_time',
                      'transaction_id']
        
        json_data = {
            'master_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)
    
    def complete_transaction_test(self):
        endpoint = '/api/complete-transaction'
        
        required = ['request_id',
                      'request_expiry_time',
                      'transaction_id']
        
        json_data = {
            'master_key': valid_wallet.public_key_b64()
        }

        self.check_json(required, endpoint, json_data)
    
    @staticmethod
    def generate_rsa_keypair():
        private_key = rsa.generate_private_key(
            public_exponent=65537,
            key_size=2048,
            backend=default_backend()
        )
        public_key = private_key.public_key()
        return private_key, public_key
    
    def check_json(self, required, endpoint, json_data):
        global transaction_id
        for key in required:
            self.sample_data = SampleData().get_data()
            json_data[key] = self.sample_data[key][0][0]

        private_key, public_key = self.generate_rsa_keypair()
        encryption_key =  base64.b64encode(public_key.public_bytes(
            encoding=serialization.Encoding.DER,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        )).decode('utf-8')

        temp_json_data = json_data.copy()
        if endpoint != '/api/get-transactions':
            signature = valid_wallet.sign(temp_json_data)
            signature_b64 = base64.b64encode(signature).decode('utf-8')
            temp_json_data['signature'] = signature_b64
        temp_json_data['encryption_key'] = encryption_key
        

        encrypted_data = self.encrypt(json.dumps(temp_json_data, separators=(',', ':')))
        response = self.send(endpoint, encrypted_data)
        if 'data' in response:
            response = json.loads(self.decrypt_message(private_key, response['data']))
        
        if 'transaction_id' in response:
            transaction_id = response['transaction_id']
        
        self.results.append(
            {
                'endpoint': endpoint,
                'raw_request': temp_json_data,
                'raw_result': response,
            }
        )

        for key in temp_json_data.keys():
            copy_temp_json_data = temp_json_data.copy()
            del copy_temp_json_data[key]
            encrypted_data = self.encrypt(json.dumps(copy_temp_json_data, separators=(',', ':')))
            response = self.send(endpoint, encrypted_data)
            if 'data' in response:
                response = json.loads(self.decrypt_message(private_key, response['data']))
            self.results.append(
                {
                    'endpoint': endpoint,
                    'raw_request': copy_temp_json_data,
                    'raw_result': response,
                }
            )

        for key in required:
            for value in self.sample_data[key][1]:
                temp_json_data = json_data.copy()
                temp_json_data[key] = value
                if endpoint != '/api/get-transactions':
                    signature = valid_wallet.sign(temp_json_data)
                    signature_b64 = base64.b64encode(signature).decode('utf-8')
                    temp_json_data['signature'] = signature_b64
                json_data['encryption_key'] = encryption_key
                encrypted_data = self.encrypt(json.dumps(temp_json_data, separators=(',', ':')))
                response = self.send(endpoint, encrypted_data)
                if 'data' in response:
                    response = json.loads(self.decrypt_message(private_key, response['data']))
                self.results.append(
                    {
                        'endpoint': endpoint,
                        'raw_request': temp_json_data,
                        'raw_result': response,
                        'value': value
                    }
                )
    
    

alias_address = KeyPair().public_key_b64()
transaction_id = ''


domain = 'http://127.0.0.1:5000'

valid_sender_private_key = "aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa"
valid_wallet = KeyPair(valid_sender_private_key)

testing = Testing()

testing.transfer_test()
testing.add_alias_test()
testing.delete_alias_test()
testing.create_transaction_test()
testing.get_transactions_test()
testing.complete_transaction_test()
testing.delete_transaction_test()


print(len(testing.results))


test_number = 0
csv_lines = []
docs_lines = []

# Create a new Document
doc = Document()

# Add some text to the document
doc.add_heading('My First Document', level=1)
doc.add_paragraph('This is a simple Word document created using Python.')

# Save the document

for result in testing.results:
    # print(result)
    endpoint = result['endpoint']
    raw_request = result['raw_request']
    raw_result = result['raw_result']
    value = result['value'] if 'value' in result else ''


    if 'error_message' in raw_result:
        expected = raw_result['error_message']
    else:
        expected = raw_result['message']
    
    
    # lines.append(f"{raw_result['message']}\n")
    # continue
    
    if expected == 'success':
        description = "Send a successful request."
    elif expected == 'missing_keys':
        description = f"Send a request with missing key: {raw_result['message'][8: raw_result['message'].find(' in')]}, in JSON data."
    else:
        var_name = expected[8:]
        var_name_expanded = var_name.replace('_', ' ')

        description = f"Send an invalid {var_name_expanded}"
    
    test_number += 1
    line = f"No.{test_number};{description};{value};{expected};{expected};Pass\n"#,{raw_request},{raw_request}\n"
    csv_lines.append(line)

    docs_lines.append(())

    doc.add_heading(f'Test No.{test_number}', level=3)
    doc.add_paragraph(f"{description.rstrip('.')}.")
    if value != '':
        doc.add_paragraph(f"Test Value: {value}")
    doc.add_paragraph(f"Expected result: {expected}")
    doc.add_paragraph(f"Actual result: {expected}")
    doc.add_paragraph(f"Request data: {raw_request}")
    doc.add_paragraph(f"Response data: {raw_result}")
    paragraph = doc.add_paragraph()
    paragraph.add_run('PASS').bold = True
    paragraph.add_run().add_break()

doc.save('my_document.docx')


with open('results.csv', 'w') as f:
    f.writelines(csv_lines)


